"""
Generates SPECYFIKACJA_BIZNESOWA.docx from scratch using python-docx.
Run: python generate_spec.py
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ────────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Helper: set paragraph font ──────────────────────────────────────────────────
def set_font(run, size=11, bold=False, color=None, italic=False):
    run.font.name  = "Calibri"
    run.font.size  = Pt(size)
    run.font.bold  = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

# ── Helper: style heading ────────────────────────────────────────────────────────
def heading(text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.name = "Calibri"
        if level == 1:
            run.font.size = Pt(18)
            run.font.color.rgb = RGBColor(0, 70, 127)
        elif level == 2:
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(31, 73, 125)
        elif level == 3:
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(54, 95, 145)
    return p

# ── Helper: normal paragraph ────────────────────────────────────────────────────
def para(text, bold=False, italic=False, size=11, color=None, space_before=0, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, italic=italic, color=color)
    return p

# ── Helper: bullet ───────────────────────────────────────────────────────────────
def bullet(text, level=0, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.left_indent  = Cm(0.5 + level * 0.5)
    if bold_prefix:
        r1 = p.add_run(bold_prefix + " ")
        set_font(r1, bold=True)
        r2 = p.add_run(text)
        set_font(r2)
    else:
        r = p.add_run(text)
        set_font(r)
    return p

# ── Helper: table ────────────────────────────────────────────────────────────────
def table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    hdr = t.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        # Blue background
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "1F497D")
        tcPr.append(shd)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(h)
        set_font(run, size=10, bold=True, color=(255, 255, 255))

    # Data rows
    for ri, row in enumerate(rows):
        tr = t.rows[ri + 1]
        bg = "F2F7FD" if ri % 2 == 0 else "FFFFFF"
        for ci, cell_text in enumerate(row):
            cell = tr.cells[ci]
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), bg)
            tcPr.append(shd)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            # Support inline bold for "**text**" pattern
            parts = cell_text.split("**")
            for idx, part in enumerate(parts):
                if part:
                    run = p.add_run(part)
                    set_font(run, size=10, bold=(idx % 2 == 1))

    # Column widths
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)

    doc.add_paragraph()  # spacer after table
    return t

# ── Helper: page break ──────────────────────────────────────────────────────────
def page_break():
    doc.add_page_break()

# ── Helper: info box (light blue) ───────────────────────────────────────────────
def info_box(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(8)
    p.paragraph_format.left_indent  = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    # Light blue shading via XML
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "DEEAF1")
    pPr.append(shd)
    run = p.add_run(text)
    set_font(run, size=10, italic=True, color=(31, 73, 125))
    return p

# ── Helper: code-style inline text ──────────────────────────────────────────────
def inline_code(p, text):
    run = p.add_run(f"`{text}`")
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(166, 28, 0)
    return run

# ═══════════════════════════════════════════════════════════════════════════════
# STRONA TYTUŁOWA
# ═══════════════════════════════════════════════════════════════════════════════

p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_title.paragraph_format.space_before = Pt(80)
run = p_title.add_run("WORKTRIPS DOC")
run.font.name = "Calibri"
run.font.size = Pt(32)
run.font.bold = True
run.font.color.rgb = RGBColor(0, 70, 127)

p_sub = doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p_sub.add_run("Specyfikacja Biznesowa Systemu")
run.font.name = "Calibri"
run.font.size = Pt(20)
run.font.color.rgb = RGBColor(54, 95, 145)

doc.add_paragraph()
doc.add_paragraph()

p_desc = doc.add_paragraph()
p_desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p_desc.add_run("System zarządzania dokumentami, partnerami i sprzedażą\ndla biur podróży")
run.font.name = "Calibri"
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(89, 89, 89)
run.font.italic = True

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

p_date = doc.add_paragraph()
p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p_date.add_run("Wersja: 2026-04-27")
run.font.name = "Calibri"
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(127, 127, 127)

page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 1. PRZEGLĄD SYSTEMU
# ═══════════════════════════════════════════════════════════════════════════════

heading("1. Przegląd systemu", 1)

info_box(
    "Worktrips Doc to wewnętrzny system SaaS obsługujący trzy główne obszary działalności biura podróży: "
    "zarządzanie dokumentami, CRM sprzedażowy oraz analitykę opartą na danych z hurtowni danych (DWH)."
)

table(
    ["Obszar", "Cel"],
    [
        ["**Zarządzanie dokumentami**", "Tworzenie, wersjonowanie, obieg podpisów i zatwierdzanie umów oraz innych dokumentów"],
        ["**CRM sprzedażowy**", "Zarządzanie leadami, konwersja na partnerów, komunikacja emailowa, aktywności handlowców"],
        ["**Analityka**", "Raporty sprzedażowe i KPI oparte na danych z DWH (systemu transakcyjnego)"],
    ],
    col_widths=[5, 11],
)

para("Logowanie:", bold=True)
para(
    "Użytkownicy logują się przez Google Workspace SAML — brak osobnych haseł w systemie. "
    "Dostęp wyłącznie dla kont w domenie @worktrips.com.",
    space_after=8,
)

heading("Moduły systemu", 2)

table(
    ["URL", "Moduł", "Dostęp"],
    [
        ["/dashboard", "Strona główna (Dashboard)", "Zalogowany"],
        ["/documents", "Zarządzanie dokumentami", "Zalogowany"],
        ["/workflow", "Workflow / Kanban zadań", "Zalogowany"],
        ["/groups", "Grupy użytkowników", "Zalogowany"],
        ["/users", "Zarządzanie użytkownikami", "Administrator"],
        ["/logs", "Logi systemowe", "Administrator"],
        ["/admin/settings", "Ustawienia aplikacji", "Administrator"],
        ["/admin/data", "Zarządzanie danymi", "Administrator"],
        ["/crm/leads", "Lista leadów sprzedażowych", "Rola CRM"],
        ["/crm/partners", "Rejestr partnerów", "Rola CRM"],
        ["/crm/onboarding", "Panel onboardingu partnerów", "Rola CRM"],
        ["/crm/reports", "Raporty sprzedaży", "Rola CRM"],
        ["/crm/calendar", "Kalendarz działań", "Rola CRM"],
        ["/crm/import", "Import CSV", "Administrator"],
        ["/crm/partner-groups", "Grupy partnerskie", "Rola CRM"],
    ],
    col_widths=[5, 8, 4],
)

page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 2. ROLE I UPRAWNIENIA
# ═══════════════════════════════════════════════════════════════════════════════

heading("2. Role i uprawnienia użytkowników", 1)

heading("2.1 Role systemowe", 2)

table(
    ["Rola", "Zakres dostępu"],
    [
        ["**Administrator**", "Pełny dostęp do wszystkich danych, ustawień i panelu administracyjnego"],
        ["**Sales Manager**", "Widzi i edytuje leady/partnerów handlowców ze swoich grup; może filtrować dane per handlowiec"],
        ["**Salesperson**", "Widzi i edytuje wyłącznie własne leady i partnerów (przypisane do niego jako assigned_to)"],
        ["**Bez roli CRM**", "Dostęp do modułu dokumentów i workflow; brak dostępu do CRM"],
    ],
    col_widths=[5, 11],
)

heading("2.2 Grupy użytkowników", 2)
bullet("Użytkownicy są przypisywani do grup (np. region, dział).")
bullet("Sales Manager widzi wszystkich handlowców z grup, do których należy.")
bullet("Grupy mogą mieć flagę ograniczenia właścicielskiego: tylko właściciel dokumentu/zasobu może go modyfikować.")
bullet("Dostęp do dokumentów grupy: wszyscy członkowie widzą dokumenty swojej grupy.")
bullet("Poziomy dostępu w grupie: read (odczyt) lub full (pełny).")
doc.add_paragraph()

heading("2.3 Zarządzanie użytkownikami", 2)
bullet("Tworzenie i dezaktywacja kont użytkowników.")
bullet("Nadawanie ról CRM (salesperson, sales_manager).")
bullet("Przypisywanie do grup z określonym poziomem dostępu.")
doc.add_paragraph()

page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 3. MODUŁ DOKUMENTÓW
# ═══════════════════════════════════════════════════════════════════════════════

heading("3. Moduł Dokumentów", 1)

heading("3.1 Typy dokumentów", 2)
table(
    ["Typ", "Opis"],
    [
        ["Umowa partnerska", "Kontrakt z partnerem handlowym"],
        ["NDA", "Umowa o zachowaniu poufności"],
        ["Umowa pracownicza", "Kontrakt z pracownikiem"],
        ["Umowa z dostawcą IT", "Kontrakt z dostawcą oprogramowania lub usług IT"],
        ["Umowa z operatorem", "Kontrakt z operatorem turystycznym"],
    ],
    col_widths=[6, 10],
)

heading("3.2 Cykl życia dokumentu", 2)
info_box(
    "new  →  being_edited  →  being_signed  →  being_approved  →  signed  →  completed\n"
    "                                                            →  rejected\n"
    "                         →  hold"
)

heading("3.3 Funkcjonalności", 2)
bullet("Tworzenie dokumentu — przypisanie do grupy, typu, stron umowy (entity1, entity2), przedmiotu, klasyfikacji GDPR.")
bullet("Wersjonowanie — każda zapisana wersja jest odrębnym rekordem; historia wszystkich wersji dostępna.")
bullet("Tagi — dowolne tagi do kategoryzacji i wyszukiwania.")
bullet("Załączniki — pliki dołączone do wersji, przechowywane w Azure Blob Storage.")
bullet("Daty — data zawarcia, data wygaśnięcia; alerty kolorowe (czerwony < 90 dni, żółty < 30 dni — konfigurowalne).")
bullet("Powiązanie z CRM — dokument może być podlinkowany do leada lub partnera.")
bullet("Obieg podpisów — integracja z zewnętrznym serwisem Signus (koperty, statusy przez webhook).")
doc.add_paragraph()

heading("3.4 Kontrola dostępu", 2)
bullet("Właściciel dokumentu.")
bullet("Przynależność do grupy dokumentu.")
bullet("Użytkownicy z aktywnym zadaniem workflow na tym dokumencie.")
bullet("Administratorzy widzą wszystko.")
doc.add_paragraph()

heading("3.5 Klasyfikacje GDPR", 2)
table(
    ["Typ", "Opis"],
    [
        ["Powierzenie przetwarzania danych", "Umowa powierzenia zgodnie z RODO Art. 28"],
        ["Administratorstwo danych", "Umowa między współadministratorami"],
        ["Bez GDPR", "Dokument nie dotyczy danych osobowych"],
    ],
    col_widths=[7, 9],
)

page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 4. WORKFLOW
# ═══════════════════════════════════════════════════════════════════════════════

heading("4. Moduł Workflow", 1)

info_box(
    "Workflow umożliwia formalny obieg dokumentów wewnątrz organizacji — "
    "przypisanie zadań edycji, podpisu, zatwierdzenia lub przeczytania konkretnym osobom."
)

heading("4.1 Typy zadań", 2)
table(
    ["Typ", "Opis"],
    [
        ["edit", "Zadanie edycji dokumentu"],
        ["sign", "Zadanie podpisania dokumentu"],
        ["approve", "Zadanie zatwierdzenia dokumentu"],
        ["read", "Zadanie zapoznania się z dokumentem"],
    ],
    col_widths=[4, 12],
)

heading("4.2 Statusy zadań", 2)
info_box("pending  →  in_progress  →  completed\n                          →  cancelled\n                          →  rejected")

heading("4.3 Widoki Workflow", 2)
bullet("Moje zadania — lista zadań przypisanych do zalogowanego użytkownika, z podziałem na statusy.")
bullet("Kanban Board — wszystkie zadania (dla managera/admina) w kolumnach według statusu.")
bullet("Auto-odświeżanie — konfigurowalne w ustawieniach (interval w sekundach).")
doc.add_paragraph()

heading("4.4 Powiadomienia", 2)
para("Przypisanie zadania generuje powiadomienie dla odbiorcy.", space_after=8)

page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 5. CRM — LEADY
# ═══════════════════════════════════════════════════════════════════════════════

heading("5. Moduł CRM — Leady sprzedażowe", 1)

info_box(
    "Lead to potencjalny partner handlowy, który jest w procesie sprzedaży. "
    "Leady prowadzone są przez handlowców i przechodzą przez kolejne etapy aż do konwersji lub utraty."
)

heading("5.1 Dane leada", 2)
table(
    ["Pole", "Opis"],
    [
        ["Firma", "Nazwa firmy potencjalnego partnera (wymagane)"],
        ["NIP", "Numer identyfikacji podatkowej"],
        ["Osoba kontaktowa", "Imię, nazwisko, stanowisko"],
        ["Email / Telefon", "Dane kontaktowe głównej osoby"],
        ["Dodatkowe kontakty", "Wiele osób kontaktowych dla jednej firmy"],
        ["Strona WWW", "Adres strony internetowej"],
        ["Branża", "Branża klienta (słownikowa)"],
        ["Źródło pozyskania", "Skąd pochodzi lead (słownikowe: strona www, polecenie, cold call, LinkedIn, targi, itp.)"],
        ["Etap sprzedaży", "Aktualny etap w lejku sprzedażowym"],
        ["Wartość kontraktu (PLN)", "Szacowana wartość umowy"],
        ["Prawdopodobieństwo (%)", "Szansa zamknięcia (0–100%)"],
        ["Planowana data zamknięcia", "Kiedy ma dojść do podpisania umowy"],
        ["Handlowiec odpowiedzialny", "Przypisany handlowiec"],
        ["Tagi", "Dowolne etykiety do kategoryzacji"],
        ["Gorący lead 🔥", "Flaga priorytetu"],
        ["Procent sprzedaży online", "Jaka część sprzedaży partnera odbywa się online (0–100%)"],
        ["Roczny obrót / waluta", "Szacowany roczny wolumen"],
        ["Data pierwszego kontaktu", "Kiedy po raz pierwszy skontaktowaliśmy się"],
        ["Notatki", "Swobodny tekst"],
        ["Powód utraty", "Wypełniany przy zamknięciu jako closed_lost"],
    ],
    col_widths=[5.5, 10.5],
)

heading("5.2 Etapy sprzedaży (pipeline)", 2)
info_box(
    "new  →  qualification  →  presentation  →  offer  →  negotiation  →  closed_won\n"
    "                                                                   →  closed_lost\n\n"
    "Po closed_won następuje konwersja na partnera."
)

heading("5.3 Widok listy leadów", 2)
bullet("Kanban — kolumny odpowiadają etapom; drag & drop do zmiany etapu.")
bullet("Filtry: szukaj po nazwie firmy, etap, źródło, handlowiec (manager), gorący, zakres dat zamknięcia.")
bullet("Statystyki nagłówkowe: liczba leadów, gorące, pipeline (wartość × prawdopodobieństwo), Won, Lost.")
bullet("Odznaki emailowe — przy każdym leadzie widoczna liczba nieprzeczytanych wiadomości.")
bullet("Auto-odświeżanie co 60 sekund.")
doc.add_paragraph()

heading("5.4 Widok szczegółów leada", 2)
para("Trzyszpaltowy układ:", bold=True)

table(
    ["Kolumna", "Zawartość"],
    [
        ["Lewa — Informacje", "Dane kontaktowe, dodatkowe kontakty, dane sprzedażowe, historia zmian"],
        ["Środkowa — Aktywności", "Chronologiczna lista aktywności (emaile, notatki, call, meeting), formularz dodawania"],
        ["Prawa — Powiązania", "Powiązane dokumenty, konto testowe, konwersja na partnera"],
    ],
    col_widths=[5, 11],
)

heading("5.5 Aktywności leada", 2)
table(
    ["Typ", "Opis"],
    [
        ["email", "Wysłana lub odebrana wiadomość email (przez Gmail)"],
        ["call", "Zarejestrowane połączenie telefoniczne z notatką i czasem trwania"],
        ["meeting", "Spotkanie z uczestnikami, lokalizacją i czasem trwania"],
        ["note", "Swobodna notatka"],
        ["doc_sent", "Wysłany dokument"],
    ],
    col_widths=[4, 12],
)

heading("5.6 Konwersja leada na partnera", 2)
bullet("Dostępna dla leadów na etapie closed_won.")
bullet("Tworzy nowy rekord w module Partnerzy na podstawie danych leada.")
bullet("Lead otrzymuje powiązanie z nowym partnerem.")
bullet("Dane (firma, kontakty, email, NIP, branża, itp.) są automatycznie przenoszone.")
doc.add_paragraph()

heading("5.7 Import leadów z CSV", 2)
bullet("Dostępny przez panel admina (/crm/import).")
bullet("Historia importów z logiem błędów i licznikami poprawnie/błędnie zaimportowanych rekordów.")
doc.add_paragraph()

page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 6. CRM — PARTNERZY
# ═══════════════════════════════════════════════════════════════════════════════

heading("6. Moduł CRM — Partnerzy", 1)

info_box(
    "Partner to firma, która podpisała umowę i korzysta z platformy Worktrips. "
    "Partnerzy mogą powstać przez konwersję leada lub być dodani bezpośrednio. "
    "Kluczowa integracja z DWH dostarcza dane sprzedażowe (tylko do odczytu)."
)

heading("6.1 Dane partnera", 2)
table(
    ["Pole", "Opis"],
    [
        ["Firma", "Nazwa firmy (wymagane)"],
        ["NIP", "Numer identyfikacji podatkowej"],
        ["Adres", "Pełny adres siedziby"],
        ["Kontakt główny", "Imię, nazwisko, stanowisko, email, telefon"],
        ["Dodatkowe kontakty", "Wiele osób kontaktowych"],
        ["Branża", "Branża (słownikowa)"],
        ["Opiekun", "Przypisany handlowiec"],
        ["Grupa partnerska", "Przynależność do grupy (regionalna/sektorowa)"],
        ["Status", "Aktualny status (onboarding, active, inactive, churned)"],
        ["Etap onboardingu", "Krok 0–3 w procesie wdrożenia"],
        ["Wartość kontraktu", "Wynegocjowana wartość umowy"],
        ["Data podpisania umowy", "Data zawarcia kontraktu"],
        ["Data wygaśnięcia umowy", "Koniec okresu umownego"],
        ["Procent sprzedaży online", "Część sprzedaży w kanale online (0–100%)"],
        ["Roczny obrót / waluta", "Wolumen obrotu partnera"],
        ["Języki obsługi", "Języki, w których partner obsługuje klientów"],
        ["Kraje operacyjne", "Kraje, w których partner działa"],
        ["Zainteresowania produktowe", "Typy produktów sprzedawane przez partnera"],
        ["Link do DWH", "Powiązanie z rekordem w hurtowni danych (read-only)"],
        ["Liczba licencji", "Liczba aktywnych licencji w systemie"],
    ],
    col_widths=[5.5, 10.5],
)

heading("6.2 Statusy partnerów", 2)
info_box("onboarding  →  active  →  inactive\n                       →  churned")

heading("6.3 Widok listy partnerów", 2)
bullet("Widok kart — kafelki z kluczowymi informacjami (firma, status, opiekun, obroty).")
bullet("Widok tabeli — szczegółowa tabela z paginacją.")
bullet("Filtry: status, opiekun, grupa, branża, szukaj po nazwie.")
bullet("Sortowanie: firma, branża, grupa, opiekun, wartość kontraktu, procent online, data wygaśnięcia.")
bullet("Odznaki emailowe — nieprzeczytane emaile przy każdym partnerze.")
bullet("Auto-odświeżanie co 60 sekund.")
doc.add_paragraph()

heading("6.4 Aktywności partnera", 2)
table(
    ["Typ", "Opis"],
    [
        ["email", "Wysłana lub odebrana wiadomość email"],
        ["call", "Połączenie telefoniczne"],
        ["meeting", "Spotkanie"],
        ["note", "Notatka"],
        ["doc_sent", "Wysłany dokument"],
        ["training", "Zarejestrowane szkolenie"],
    ],
    col_widths=[4, 12],
)

heading("6.5 Grupy partnerskie", 2)
bullet("Grupy służą do organizacji portfela partnerów (np. region, sektor).")
bullet("Każda grupa ma opiekuna-managera.")
bullet("Partnerzy przypisani do jednej grupy.")
bullet("Panel /crm/partner-groups — tworzenie i edycja grup.")
doc.add_paragraph()

page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 7. ONBOARDING PARTNERÓW
# ═══════════════════════════════════════════════════════════════════════════════

heading("7. Moduł CRM — Onboarding partnerów", 1)

info_box(
    "Onboarding to 4-etapowy ustrukturyzowany proces wprowadzenia nowego partnera do platformy Worktrips. "
    "Każdy etap ma przypisane zadania dla handlowców."
)

heading("7.1 Etapy onboardingu", 2)
table(
    ["Krok", "Nazwa", "Opis"],
    [
        ["0", "Umowa", "Formalizacja umowy partnerskiej"],
        ["1", "Konfiguracja", "Konfiguracja konta i dostępów w systemie"],
        ["2", "Szkolenie", "Szkolenie użytkowników partnera"],
        ["3", "Uruchomienie", "Gotowy do aktywnej sprzedaży"],
    ],
    col_widths=[2, 4, 10],
)

heading("7.2 Zadania onboardingowe", 2)
table(
    ["Pole", "Opis"],
    [
        ["Krok", "Przynależność do kroku 0–3"],
        ["Typ", "task, call, email, meeting, note, doc_sent, training"],
        ["Tytuł i opis", "Treść zadania"],
        ["Przypisana osoba", "Handlowiec odpowiedzialny za wykonanie"],
        ["Termin", "Data wykonania (due_date)"],
        ["Status", "Otwarte / Zakończone (z datą i osobą zamykającą)"],
    ],
    col_widths=[5, 11],
)

heading("7.3 Szablony zadań", 2)
bullet("Administrator definiuje szablony zadań dla każdego kroku onboardingu.")
bullet("Przy tworzeniu onboardingu szablony są automatycznie przekształcane w zadania.")
bullet("Szablon zawiera: krok, typ, tytuł, treść, liczbę dni od startu (termin relatywny).")
doc.add_paragraph()

heading("7.4 Widoki panelu onboardingu", 2)
table(
    ["Widok", "Opis"],
    [
        ["Partners", "Lista partnerów aktualnie w onboardingu z postępem na każdym kroku"],
        ["Kanban", "Partnerzy w kolumnach według kroku; przeciąganie do zmiany kroku"],
        ["Timeline", "Oś czasowa postępu wszystkich partnerów"],
        ["Calendar", "Widok kalendarzowy terminów zadań"],
    ],
    col_widths=[4, 12],
)

page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 8. GMAIL / EMAIL
# ═══════════════════════════════════════════════════════════════════════════════

heading("8. Moduł CRM — Komunikacja emailowa (Gmail)", 1)

heading("8.1 Podłączenie konta Gmail", 2)
bullet("Każdy handlowiec łączy własne konto Gmail przez OAuth2 (Google).")
bullet("Połączenie odbywa się jednorazowo z poziomu modalu Email w CRM.")
bullet("Po połączeniu Gmail jest używany do wysyłki w imieniu handlowca.")
bullet("Możliwość ponownej autoryzacji (reauth) z poziomu UI przy wygaśnięciu tokenów.")
doc.add_paragraph()

heading("8.2 Okno korespondencji (Compose)", 2)
table(
    ["Element formularza", "Opis"],
    [
        ["Do", "Autocomplete z kontaktów leada/partnera; możliwość wpisania dowolnego adresu"],
        ["DW (CC)", "Kopie wiadomości"],
        ["Temat", "Temat wiadomości"],
        ["Treść", "Tekst z opcją dołączenia historii korespondencji jako cytat"],
        ["Załączniki — dysk lokalny", "Upload plików z komputera (maks. 25 MB/plik)"],
        ["Załączniki — Google Drive", "Picker wyboru pliku z Drive (przycisk 📁 Z Google Drive); Google Docs/Sheets/Slides/Drawings eksportowane do PDF"],
    ],
    col_widths=[6, 10],
)

heading("8.3 Wątki emailowe (Thread Viewer)", 2)
bullet("Z panelu aktywności można otworzyć wątek emailowy — pełna historia konwersacji w formacie Gmail.")
bullet("Każda wiadomość: nadawca, odbiorcy, data, treść, załączniki.")
bullet("Nieprzeczytane wiadomości od klienta: żółte tło + pogrubiony tekst.")
bullet("Pobieranie i podgląd załączników (pliki otwierane w przeglądarce).")
bullet("Przycisk Odpowiedz — formularz odpowiedzi w tym samym oknie (treść, załączniki lokalne + Drive).")
doc.add_paragraph()

heading("8.4 Odbieranie emaili (automatyczne)", 2)
bullet("System odbiera wiadomości przychodzące przez Google Cloud Pub/Sub w czasie rzeczywistym.")
bullet("Każdy email dopasowywany do leada/partnera na podstawie adresów nadawcy/odbiorcy.")
bullet("Dopasowany email tworzy aktywność email z flagą nieprzeczytana.")
bullet("Odznaki nieprzeczytanych emaili (badge z liczbą) widoczne na liście leadów/partnerów i w szczegółach.")
bullet("Odczytanie wątku automatycznie oznacza wiadomości jako przeczytane.")
bullet("Nowe adresy emailowe kontaktów są automatycznie zapisywane do bazy kontaktów.")
doc.add_paragraph()

heading("8.5 Załączniki emailowe", 2)
bullet("Załączniki z odebranych emaili rejestrowane i przechowywane w Azure Blob Storage.")
bullet("Możliwość podglądu (w przeglądarce) i pobrania.")
bullet("Wysłane załączniki widoczne przy wiadomościach wychodzących.")
doc.add_paragraph()

page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 9. RAPORTY I ANALITYKA
# ═══════════════════════════════════════════════════════════════════════════════

heading("9. Moduł CRM — Raporty i analityka sprzedaży", 1)

heading("9.1 Dashboard KPI (/crm/reports)", 2)
table(
    ["Wskaźnik (KPI)", "Opis"],
    [
        ["Pipeline", "Suma wartości aktywnych leadów × prawdopodobieństwo"],
        ["Won", "Wartość zamkniętych wygranych w wybranym okresie"],
        ["Win Rate", "Procent leadów zamkniętych jako wygrane"],
        ["Avg Cycle", "Średni czas domknięcia transakcji (w dniach)"],
        ["Active Leads", "Liczba aktywnych leadów"],
        ["Hot Leads", "Liczba leadów oznaczonych jako gorące"],
    ],
    col_widths=[5, 11],
)

heading("9.2 Wykresy i zestawienia", 2)
table(
    ["Wykres / Zestawienie", "Opis"],
    [
        ["Lejek sprzedaży", "Rozkład leadów po etapach"],
        ["Monthly Trend", "Porównanie pipeline vs. won w kolejnych miesiącach"],
        ["Sales by Person", "Tabela per handlowiec: leady, pipeline, won, win%, pasek postępu"],
        ["Source Breakdown", "Rozkład źródeł pozyskania leadów"],
        ["Lost Reasons", "Analiza powodów utraty leadów"],
        ["Velocity by Stage", "Średni czas spędzony przez leada na każdym etapie"],
    ],
    col_widths=[6, 10],
)

para("Filtry raportów:", bold=True)
bullet("Okres: Q1, Q2, YTD, roczny, niestandardowy")
bullet("Handlowiec / Sales Manager")
bullet("Partner (powiązany z DWH)")
bullet("Kategoria produktu")
para("Eksport: PDF", space_after=8)

heading("9.3 Dane sprzedażowe z DWH", 2)
para("Dostępne agregacje per partner:", bold=False)
bullet("Łączny obrót brutto / netto (PLN)")
bullet("Prowizje i marże")
bullet("Liczba transakcji")
bullet("Liczba pasażerów (pax)")
bullet("Podział po kategoriach produktowych:")
bullet("hotel, lot, pociąg, autobus, prom, wynajem auta, transfer, ubezpieczenie turystyczne, wiza, inne", level=1)
doc.add_paragraph()

heading("9.4 Budżety sprzedażowe", 2)
bullet("Definiowanie planów sprzedażowych per handlowiec / miesiąc.")
bullet("Porównanie planu vs. wykonania w raportach.")
doc.add_paragraph()

page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 10. KALENDARZ
# ═══════════════════════════════════════════════════════════════════════════════

heading("10. Moduł CRM — Kalendarz", 1)

bullet("Wyświetlanie wszystkich zaplanowanych aktywności handlowców (call, meeting) w widoku kalendarza.")
bullet("Filtrowanie per handlowiec.")
bullet("Integracja z Google Calendar — tworzenie eventów przez Service Account (Domain-Wide Delegation).")
doc.add_paragraph()

page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 11. PANEL ADMINISTRATORA
# ═══════════════════════════════════════════════════════════════════════════════

heading("11. Panel Administratora", 1)

heading("11.1 Ustawienia aplikacji (/admin/settings)", 2)
para(
    "Edycja parametrów konfiguracyjnych systemu z poziomu UI bez konieczności zmiany kodu. "
    "Każde ustawienie ma: klucz, wartość, typ (number/boolean/string/json), kategorię i opis.",
    space_after=8,
)
table(
    ["Klucz ustawienia", "Opis", "Domyślnie"],
    [
        ["expiration_red_days", "Dni do wygaśnięcia — alarm czerwony", "90"],
        ["expiration_soon_days", "Dni do wygaśnięcia — ostrzeżenie żółte", "30"],
        ["kanban_refresh_interval_sec", "Auto-odświeżanie Kanban (0 = wyłączone)", "0"],
        ["default_page_size", "Domyślna paginacja list", "50"],
        ["lead_attachments_folder_url", "URL folderu Google Drive dla załączników", "—"],
    ],
    col_widths=[6, 8, 3],
)

heading("11.2 Logi systemowe (/logs)", 2)
para("Historia zdarzeń systemowych, błędów i operacji użytkowników (audit trail). Dostęp tylko dla administratora.", space_after=8)

heading("11.3 Zarządzanie użytkownikami (/users)", 2)
bullet("Przeglądanie listy użytkowników (z datasource SAML).")
bullet("Nadawanie / odbieranie ról CRM.")
bullet("Aktywacja / dezaktywacja kont.")
bullet("Przypisywanie do grup z poziomem dostępu.")
doc.add_paragraph()

heading("11.4 Zarządzanie grupami (/groups)", 2)
bullet("Tworzenie i edycja grup użytkowników.")
bullet("Przypisywanie członków i ról w grupie.")
bullet("Konfiguracja ograniczeń właścicielskich.")
doc.add_paragraph()

heading("11.5 Import CRM (/crm/import)", 2)
bullet("Import leadów z pliku CSV.")
bullet("Podgląd historii importów z logiem błędów i licznikami.")
doc.add_paragraph()

heading("11.6 Zarządzanie danymi (/admin/data)", 2)
para("Narzędzia administracyjne do zarządzania danymi aplikacji — backup, reset, czyszczenie danych testowych.", space_after=8)

page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 12. INTEGRACJE ZEWNĘTRZNE
# ═══════════════════════════════════════════════════════════════════════════════

heading("12. Integracje zewnętrzne", 1)

heading("12.1 Google Workspace", 2)
table(
    ["Integracja", "Użycie w systemie"],
    [
        ["SAML SSO", "Logowanie użytkowników przez konta Google Workspace (brak haseł w systemie)"],
        ["Gmail OAuth2", "Wysyłanie i odbieranie emaili w imieniu każdego handlowca (per user OAuth)"],
        ["Google Cloud Pub/Sub", "Powiadomienia o nowych emailach w czasie rzeczywistym (polling co 30s)"],
        ["Google Drive Picker", "Wybór plików z Drive do załączenia do emaila (browser-side GAPI)"],
        ["Google Calendar API", "Tworzenie eventów kalendarza przez Service Account z delegacją domeny"],
    ],
    col_widths=[5, 11],
)

heading("12.2 Azure Blob Storage", 2)
bullet("Przechowywanie załączników emailowych (wysłanych i odebranych).")
bullet("Przechowywanie wersji dokumentów.")
bullet("Trwały, skalowalny magazyn plików binarnych.")
doc.add_paragraph()

heading("12.3 Signus (podpisy elektroniczne)", 2)
bullet("Tworzenie kopert do podpisu elektronicznego na podstawie dokumentów z systemu.")
bullet("Webhook odbierający aktualizacje statusu podpisu od Signus.")
bullet("Automatyczna zmiana statusu dokumentu po podpisaniu przez wszystkie strony.")
doc.add_paragraph()

heading("12.4 Platform API (system transakcyjny Worktrips)", 2)
bullet("Zewnętrzna platforma Worktrips pushuje transakcje do CRM przez REST API (klucz API w nagłówku).")
bullet("Dane transakcji: numer rezerwacji, data, wartość netto/brutto, prowizja, marża, produkty, pasażerowie.")
bullet("Transakcje służą jako dane operacyjne uzupełniające DWH.")
doc.add_paragraph()

heading("12.5 DWH — Hurtownia danych", 2)
bullet("Schemat dwh w tej samej bazie PostgreSQL, zasilany przez zewnętrzny proces ETL.")
bullet("Tabele: dwh.dm_partner (wymiar partnerów), dwh.dm_sales (fakty sprzedażowe).")
bullet("Dane z DWH są tylko do odczytu w CRM — służą do wyświetlania historii i raportów.")
bullet("Partnerzy CRM powiązani z rekordami DWH przez unikalny identyfikator dwh_partner_id.")
doc.add_paragraph()

page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 13. SŁOWNIKI I KONFIGURACJA
# ═══════════════════════════════════════════════════════════════════════════════

heading("13. Słowniki i konfiguracja", 1)

para(
    "Wszystkie poniższe wartości są konfigurowane przez administratora w panelu ustawień. "
    "Zmiana słownika natychmiast wpływa na dostępne opcje w formularzach.",
    space_after=8,
)

table(
    ["Słownik", "Wartości"],
    [
        ["Źródła leadów", "strona www, polecenie, cold call, LinkedIn, targi, partner, agent, kampania, inbound, inne"],
        ["Etapy leadów", "new, qualification, presentation, offer, negotiation, closed_won, closed_lost"],
        ["Statusy partnerów", "onboarding, active, inactive, churned"],
        ["Stanowiska kontaktów", "CEO, CFO, CTO, COO, VP, Dyrektor, Manager, Specjalista, Właściciel, Inne"],
        ["Branże", "IT, Finanse, Transport, Turystyka, Zdrowie, Handel, Produkcja, Prawnicza, Edukacja, Inne"],
        ["Waluty", "PLN, EUR, USD, GBP, CHF"],
        ["Kategorie produktów", "hotel, lot, pociąg, autobus, prom, wynajem auta, transfer, ubezpieczenie, wiza, inne"],
        ["Typy dokumentów", "partner_agreement, nda, employee_agreement, it_supplier_agreement, operator_agreement"],
        ["Klasyfikacje GDPR", "Powierzenie przetwarzania danych, Administratorstwo danych, Bez GDPR"],
        ["Szablony onboardingu", "JSON: krok, typ zadania, tytuł, treść, dni od startu"],
    ],
    col_widths=[5, 11],
)

page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 14. MODEL DANYCH
# ═══════════════════════════════════════════════════════════════════════════════

heading("14. Model danych — tabele główne", 1)

heading("14.1 Tabele dokumentów", 2)
table(
    ["Tabela", "Opis"],
    [
        ["documents", "Dokumenty — główny rejestr"],
        ["document_versions", "Wersje dokumentów (pełna historia)"],
        ["document_tags", "Tagi dokumentów"],
        ["workflow_tasks", "Zadania obiegu dokumentów"],
        ["attachments", "Pliki załączone do wersji dokumentu"],
    ],
    col_widths=[6, 10],
)

heading("14.2 Tabele CRM — Leady", 2)
table(
    ["Tabela", "Opis"],
    [
        ["crm_leads", "Leady sprzedażowe"],
        ["crm_lead_activities", "Aktywności (email, call, meeting, note, doc_sent)"],
        ["crm_lead_contacts", "Dodatkowe kontakty leada"],
        ["crm_lead_documents", "Powiązania lead ↔ dokument"],
    ],
    col_widths=[6, 10],
)

heading("14.3 Tabele CRM — Partnerzy", 2)
table(
    ["Tabela", "Opis"],
    [
        ["crm_partners", "Partnerzy handlowi"],
        ["crm_partner_groups", "Grupy partnerów"],
        ["crm_partner_activities", "Aktywności partnera"],
        ["crm_partner_contacts", "Dodatkowe kontakty partnera"],
        ["crm_onboarding_tasks", "Zadania procesu onboardingu"],
        ["crm_onboarding_templates", "Szablony zadań onboardingowych"],
    ],
    col_widths=[6, 10],
)

heading("14.4 Tabele email i transakcje", 2)
table(
    ["Tabela", "Opis"],
    [
        ["user_gmail_tokens", "Tokeny OAuth Gmail per użytkownik"],
        ["crm_email_attachments", "Załączniki emailowe (metadata + ścieżka w Blob Storage)"],
        ["crm_email_message_reads", "Statusy przeczytania wiadomości email"],
        ["crm_transactions", "Transakcje z platformy Worktrips"],
        ["crm_transaction_products", "Produkty w transakcji"],
    ],
    col_widths=[6, 10],
)

heading("14.5 Tabele DWH", 2)
table(
    ["Tabela", "Opis"],
    [
        ["dwh.dm_partner", "Wymiar partnerów z systemu transakcyjnego (ETL)"],
        ["dwh.dm_sales", "Fakty sprzedażowe (ETL — obroty, prowizje, marże)"],
    ],
    col_widths=[6, 10],
)

heading("14.6 Tabele systemowe", 2)
table(
    ["Tabela", "Opis"],
    [
        ["users", "Konta użytkowników (z SAML)"],
        ["group_profiles", "Grupy użytkowników"],
        ["user_group_roles", "Przypisania użytkowników do grup z rolami"],
        ["app_settings", "Konfiguracja aplikacji (klucz-wartość)"],
        ["crm_sales_budgets", "Budżety sprzedażowe per handlowiec/miesiąc"],
        ["crm_import_logs", "Historia importów CSV"],
        ["audit_logs", "Logi audytu (historia operacji użytkowników)"],
    ],
    col_widths=[6, 10],
)

# ── Final note ──────────────────────────────────────────────────────────────────
doc.add_paragraph()
doc.add_paragraph()
p_footer = doc.add_paragraph()
p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p_footer.add_run("Dokument wygenerowany automatycznie na podstawie analizy kodu źródłowego. Wersja 2026-04-27.")
set_font(run, size=9, italic=True, color=(127, 127, 127))

# ── Anonymisation: replace brand/partner terms ───────────────────────────────────
_REPL = [
    # App name first (longer match wins)
    ("WORKTRIPS DOC", "x-CRM"), ("Worktrips Doc", "x-CRM"),
    # Company name
    ("@worktrips.com", "@firma.com"), ("Worktrips", "Firma"), ("worktrips", "firma"),
    # Partner (pl) – always longer variants first
    ("partnerskich", "klienckich"), ("partnerskim", "klienckim"),
    ("partnerskie", "klienckie"), ("partnerską", "kliencką"),
    ("partnerska", "kliencka"), ("partnerski", "kliencki"),
    ("partnerów", "klientów"), ("partnerom", "klientom"),
    ("partnerami", "klientami"), ("partnerach", "klientach"),
    ("partnerze", "kliencie"), ("partnera", "klienta"),
    ("partnerowi", "klientowi"), ("partnerem", "klientem"),
    ("partnerzy", "klienci"), ("partner", "klient"),
    # Capitalised
    ("Partnerskich", "Klienckich"), ("Partnerskim", "Klienckim"),
    ("Partnerskie", "Klienckie"), ("Partnerską", "Kliencką"),
    ("Partnerska", "Kliencka"), ("Partnerski", "Kliencki"),
    ("Partnerów", "Klientów"), ("Partnerom", "Klientom"),
    ("Partnerami", "Klientami"), ("Partnerach", "Klientach"),
    ("Partnerze", "Kliencie"), ("Partnera", "Klienta"),
    ("Partnerowi", "Klientowi"), ("Partnerem", "Klientem"),
    ("Partnerzy", "Klienci"), ("Partner", "Klient"),
    ("PARTNER", "KLIENT"),
]

def _anon(text):
    for old, new in _REPL:
        text = text.replace(old, new)
    return text

def _anon_para(para):
    for run in para.runs:
        run.text = _anon(run.text)

for _p in doc.paragraphs:
    _anon_para(_p)
for _t in doc.tables:
    for _r in _t.rows:
        for _c in _r.cells:
            for _p in _c.paragraphs:
                _anon_para(_p)

# ── Save ─────────────────────────────────────────────────────────────────────────
output_path = r"C:\Users\Adam\Documents\worktrips-doc-frontend\xCRM_Specyfikacja_Biznesowa.docx"
doc.save(output_path)
print(f"Zapisano: {output_path}")
